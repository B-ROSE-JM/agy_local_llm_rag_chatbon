"""
문서 로더: Word (.docx) 및 PDF 파일을 페이지 단위로 파싱합니다.

각 페이지는 DocumentPage 데이터 구조로 반환되며,
파일명, 파일 경로, 페이지 번호, 텍스트, 메타데이터를 포함합니다.
"""
import os
from pathlib import Path
from dataclasses import dataclass, field
from typing import List, Optional, Dict, Any
from src.utils.logging_utils import get_logger

logger = get_logger("document_loader")


@dataclass
class DocumentPage:
    """파싱된 문서의 개별 페이지를 나타냅니다."""
    file_name: str
    file_path: str
    page_number: int
    text: str
    total_pages: int
    metadata: Dict[str, Any] = field(default_factory=dict)


def load_document(file_path: str) -> List[DocumentPage]:
    """
    단일 문서 파일을 로드하여 페이지 목록으로 반환합니다.

    Args:
        file_path: 문서 파일의 절대 경로

    Returns:
        List[DocumentPage]: 페이지별 파싱 결과

    Raises:
        FileNotFoundError: 파일이 존재하지 않는 경우
        ValueError: 지원하지 않는 파일 형식인 경우
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"파일을 찾을 수 없습니다: {file_path}")

    ext = path.suffix.lower()
    if ext == ".pdf":
        return _load_pdf(path)
    elif ext == ".docx":
        return _load_docx(path)
    else:
        raise ValueError(
            f"지원하지 않는 파일 형식입니다: {ext}. "
            f"지원 형식: .docx, .pdf"
        )


def load_documents_from_directory(
    directory: str,
    extensions: Optional[set] = None,
) -> List[DocumentPage]:
    """
    디렉토리 내 모든 지원 문서를 파싱합니다.

    Args:
        directory: 문서가 저장된 디렉토리 경로
        extensions: 파싱할 파일 확장자 집합 (기본: .docx, .pdf)

    Returns:
        모든 문서의 페이지 목록
    """
    if extensions is None:
        extensions = {".docx", ".pdf"}

    dir_path = Path(directory)
    if not dir_path.exists():
        raise FileNotFoundError(f"디렉토리를 찾을 수 없습니다: {directory}")

    all_pages: List[DocumentPage] = []
    files = sorted(
        f for f in dir_path.iterdir()
        if f.is_file() and f.suffix.lower() in extensions
    )

    logger.info(f"디렉토리 '{directory}'에서 {len(files)}개 문서 발견")

    for file_path in files:
        try:
            pages = load_document(str(file_path))
            all_pages.extend(pages)
            logger.info(f"  ✓ {file_path.name}: {len(pages)} 페이지 파싱 완료")
        except Exception as e:
            logger.error(f"  ✗ {file_path.name} 파싱 실패: {e}")

    logger.info(f"총 {len(all_pages)} 페이지 파싱 완료 ({len(files)}개 파일)")
    return all_pages


def _load_pdf(path: Path) -> List[DocumentPage]:
    """
    PyMuPDF(fitz)를 사용하여 PDF를 페이지 단위로 파싱합니다.
    정확한 페이지 번호를 보존합니다.
    """
    import fitz  # PyMuPDF

    pages: List[DocumentPage] = []
    file_name = path.name
    file_path_str = str(path.resolve())

    try:
        doc = fitz.open(str(path))
        total_pages = len(doc)

        for page_num in range(total_pages):
            page = doc[page_num]
            text = page.get_text("text")

            # 빈 페이지 스킵 (이미지만 있는 페이지 등)
            cleaned_text = text.strip()
            if not cleaned_text:
                logger.debug(
                    f"  {file_name} 페이지 {page_num + 1}: 텍스트 없음 (스킵)"
                )
                continue

            pages.append(DocumentPage(
                file_name=file_name,
                file_path=file_path_str,
                page_number=page_num + 1,  # 1-indexed
                text=cleaned_text,
                total_pages=total_pages,
                metadata={
                    "format": "pdf",
                    "width": page.rect.width,
                    "height": page.rect.height,
                },
            ))

        doc.close()
        logger.info(
            f"PDF 파싱 완료: {file_name} "
            f"({total_pages}페이지 중 {len(pages)}페이지 텍스트 추출)"
        )

    except Exception as e:
        logger.error(f"PDF 파싱 오류 ({file_name}): {e}")
        raise

    return pages


def _load_docx(path: Path) -> List[DocumentPage]:
    """
    python-docx를 사용하여 Word 문서를 파싱합니다.

    Word 문서는 명시적 페이지 구분이 없으므로,
    페이지 브레이크(page break)와 단락 길이 기반으로 페이지를 추정합니다.
    """
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    pages: List[DocumentPage] = []
    file_name = path.name
    file_path_str = str(path.resolve())

    try:
        doc = Document(str(path))

        # 문서 메타데이터 추출
        doc_metadata = {}
        if doc.core_properties:
            props = doc.core_properties
            if props.title:
                doc_metadata["doc_title"] = props.title
            if props.author:
                doc_metadata["doc_author"] = props.author
            if props.created:
                doc_metadata["doc_created"] = str(props.created)

        # 단락별로 텍스트를 추출하고, 페이지 브레이크를 기준으로 분할
        current_page_texts: List[str] = []
        page_number = 1

        # 표 텍스트도 추출
        table_texts: Dict[int, List[str]] = {}  # paragraph_index -> table texts
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        # 표 텍스트는 해당 위치의 가장 가까운 단락에 연결
                        if page_number not in table_texts:
                            table_texts[page_number] = []
                        table_texts[page_number].append(cell_text)

        for para in doc.paragraphs:
            # 페이지 브레이크 감지
            has_page_break = False
            for run in para.runs:
                if run._element.xml and "w:br" in run._element.xml:
                    if 'w:type="page"' in run._element.xml:
                        has_page_break = True
                        break

            if has_page_break and current_page_texts:
                # 현재 페이지 저장
                page_text = "\n".join(current_page_texts).strip()
                if page_text:
                    meta = {"format": "docx"}
                    meta.update(doc_metadata)
                    pages.append(DocumentPage(
                        file_name=file_name,
                        file_path=file_path_str,
                        page_number=page_number,
                        text=page_text,
                        total_pages=0,  # 나중에 업데이트
                        metadata=meta,
                    ))
                page_number += 1
                current_page_texts = []

            para_text = para.text.strip()
            if para_text:
                # 헤딩 정보 포함
                if para.style and para.style.name.startswith("Heading"):
                    heading_level = para.style.name.replace("Heading", "").strip()
                    para_text = f"[제목 {heading_level}] {para_text}"
                current_page_texts.append(para_text)

        # 마지막 페이지 저장
        if current_page_texts:
            page_text = "\n".join(current_page_texts).strip()
            if page_text:
                meta = {"format": "docx"}
                meta.update(doc_metadata)
                pages.append(DocumentPage(
                    file_name=file_name,
                    file_path=file_path_str,
                    page_number=page_number,
                    text=page_text,
                    total_pages=0,
                    metadata=meta,
                ))

        # 페이지 브레이크가 전혀 없는 경우, 긴 문서를 일정 길이로 분할
        if len(pages) <= 1 and pages:
            full_text = pages[0].text
            # 약 3000자 단위로 가상 페이지 분할 (A4 1페이지 ≈ 1500~3000자)
            CHARS_PER_PAGE = 2500
            if len(full_text) > CHARS_PER_PAGE * 1.5:
                pages = []
                paragraphs = full_text.split("\n")
                current_chars = 0
                current_texts = []
                vpage = 1

                for para in paragraphs:
                    current_texts.append(para)
                    current_chars += len(para) + 1
                    if current_chars >= CHARS_PER_PAGE:
                        meta = {"format": "docx", "virtual_page": True}
                        meta.update(doc_metadata)
                        pages.append(DocumentPage(
                            file_name=file_name,
                            file_path=file_path_str,
                            page_number=vpage,
                            text="\n".join(current_texts).strip(),
                            total_pages=0,
                            metadata=meta,
                        ))
                        vpage += 1
                        current_chars = 0
                        current_texts = []

                if current_texts:
                    meta = {"format": "docx", "virtual_page": True}
                    meta.update(doc_metadata)
                    pages.append(DocumentPage(
                        file_name=file_name,
                        file_path=file_path_str,
                        page_number=vpage,
                        text="\n".join(current_texts).strip(),
                        total_pages=0,
                        metadata=meta,
                    ))

        # total_pages 업데이트
        total = len(pages)
        for p in pages:
            p.total_pages = total

        logger.info(f"DOCX 파싱 완료: {file_name} ({total} 페이지)")

    except Exception as e:
        logger.error(f"DOCX 파싱 오류 ({file_name}): {e}")
        raise

    return pages
