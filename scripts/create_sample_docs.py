import os
import sys
from pathlib import Path

# 프로젝트 루트를 sys.path에 추가하여 config 등을 활용할 수 있게 함
PROJECT_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from config import RESEARCH_NOTES_DIR

# 디렉토리 생성
RESEARCH_NOTES_DIR.mkdir(parents=True, exist_ok=True)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: python-docx가 설치되어 있지 않습니다. pip install python-docx를 실행하세요.")
    sys.exit(1)


def create_battery_note():
    """연구노트 1: 전고체 배터리 고체전해질 최적화"""
    doc = Document()
    
    # 스타일 설정
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)
    
    # 제목
    h = doc.add_heading("차세대 고에너지 밀도 전고체 배터리용 황화물계 고체전해질 합성 공정 최적화", level=1)
    h.runs[0].font.name = '맑은 고딕'
    h.runs[0].font.size = Pt(18)
    h.runs[0].font.bold = True
    
    # 메타데이터
    p_meta = doc.add_paragraph()
    p_meta.add_run("작성자: ").bold = True
    p_meta.add_run("김민우 책임연구원 (차세대전지 연구그룹)\n")
    p_meta.add_run("작성일: ").bold = True
    p_meta.add_run("2026-05-15\n")
    p_meta.add_run("프로젝트코드: ").bold = True
    p_meta.add_run("SSB-2026-04\n")
    p_meta.add_run("문서보안등급: ").bold = True
    p_meta.add_run("대외비 (Confidential)")
    
    # 1페이지: 서론 및 연구 목적
    doc.add_heading("1. 서론 및 연구 목적", level=2)
    doc.add_paragraph(
        "리튬 이온 전지의 가연성 유기 액체 전해질 사용으로 인한 화재 및 폭발 위험성을 원천적으로 배제하기 위해, "
        "비가연성 고체 전해질을 사용하는 전고체 배터리 연구가 전 세계적으로 진행 중이다. 고체 전해질 중에서 "
        "황화물계 고체 전해질은 산화물계나 고분자계에 비해 실온에서 높은 이온 전도도와 우수한 기계적 연성을 "
        "지녀 가장 유망한 후보 물질로 꼽힌다."
    )
    doc.add_paragraph(
        "본 연구의 목적은 아지로다이트(Argyrodite) 구조를 가진 황화물계 고체 전해질인 Li6PS5Cl의 볼밀링 및 소결 합성 공정을 "
        "최적화하여, 상온(25°C)에서 이온 전도도 3.5 mS/cm 이상을 달성하고 최적의 열처리 온도를 도출하는 것이다."
    )
    
    doc.add_page_break()
    
    # 2페이지: 실험 방법
    doc.add_heading("2. 실험 방법", level=2)
    doc.add_paragraph(
        "실험은 다음과 같은 화학양론적 비율에 따라 출발 원료를 혼합하여 수행되었다. "
        "Li2S (순도 99.9%), P2S5 (순도 99%), LiCl (순도 99%)를 몰비 기준으로 정밀 계량하여 Li6PS5Cl 조성을 설계하였다."
    )
    doc.add_paragraph("기계적 합금화(Mechanical Alloying)를 위해 유성볼밀(Planetary Ball Mill)을 사용하였으며, 상세 조건은 다음과 같다:")
    
    doc.add_paragraph("• 회전 속도: 450 rpm", style='List Bullet')
    doc.add_paragraph("• 밀링 시간: 12 시간 (2시간 작동 후 15분 휴지 반복)", style='List Bullet')
    doc.add_paragraph("• 용기 및 볼: 지르코니아(ZrO2) 재질 사용 (볼-대-분말 질량비 = 15:1)", style='List Bullet')
    
    doc.add_paragraph(
        "이후 얻어진 비정질 분말을 직경 13 mm 펠렛 형태로 200 MPa의 압력으로 가압 성형한 뒤, "
        "석영관에 장입하여 진공(10^-3 Torr 이하) 밀봉하였다. 열처리 온도는 500°C, 550°C, 600°C의 세 가지 조건으로 "
        "설정하였으며, 승온 속도는 분당 5°C로 하여 아르곤 분위기 하에서 6시간 동안 소결을 진행한 후 자연 냉각시켰다."
    )
    
    doc.add_page_break()
    
    # 3페이지: 실험 결과 및 분석
    doc.add_heading("3. 실험 결과 및 분석", level=2)
    doc.add_paragraph(
        "합성된 Li6PS5Cl 고체전해질의 결정 구조 분석을 위해 X선 회절분석(XRD)을 수행하였다. 분석 결과, "
        "550°C에서 소결한 샘플은 미반응 잔류물인 Li2S 피크가 완전히 소멸하고 순수한 아지로다이트 단일상이 형성되었음을 확인하였다."
    )
    doc.add_paragraph(
        "상온(25°C)에서 전기화학적 임피던스 분광법(EIS)을 사용하여 주파수 범위 1 Hz ~ 1 MHz 조건으로 이온 전도도를 "
        "측정한 결과는 아래 [표 1]과 같다."
    )
    
    # 표 1 추가
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    # 헤더
    hdr = table.rows[0].cells
    hdr[0].text = '열처리 온도 (°C)'
    hdr[1].text = '이온 전도도 (mS/cm)'
    hdr[2].text = '활성화 에너지 (eV)'
    
    # 데이터 행 1
    row1 = table.rows[1].cells
    row1[0].text = '500°C'
    row1[1].text = '2.10'
    row1[2].text = '0.25'
    
    # 데이터 행 2
    row2 = table.rows[2].cells
    row2[0].text = '550°C (최적조건)'
    row2[1].text = '3.82'
    row2[2].text = '0.22'
    
    # 데이터 행 3
    row3 = table.rows[3].cells
    row3[0].text = '600°C'
    row3[1].text = '1.50'
    row3[2].text = '0.28'
    
    doc.add_paragraph(
        "\n분석 결과, 550°C에서 소결한 고체 전해질 샘플이 최고 이온 전도도인 3.82 mS/cm를 나타냈으며, "
        "활성화 에너지는 0.22 eV로 가장 낮았다. 반면 600°C에서는 Li6PS5Cl 분말의 고온 분해 반응이 일어나 전도도가 급격히 저하되었다."
    )
    
    doc.add_page_break()
    
    # 4페이지: 결론 및 향후 계획
    doc.add_heading("4. 결론 및 향후 계획", level=2)
    doc.add_paragraph(
        "본 연구를 통해 유성볼밀 공정과 550°C 열처리 공정의 최적 조합을 도출하여 이온 전도도 3.82 mS/cm의 "
        "황화물계 고체전해질을 합성하는 데 성공하였다. 이는 당사 개발 목표치인 3.5 mS/cm를 크게 상회하는 성과이다."
    )
    doc.add_paragraph("향후 연구 계획은 다음과 같다:")
    doc.add_paragraph("• 1단계: NCM811 하이니켈 양극 및 실리콘 음극을 결합한 전고체 전지 셀 제작", style='List Bullet')
    doc.add_paragraph("• 2단계: 상온 및 고온(45°C)에서의 충방전 사이클 수명(100회) 평가 및 용량 유지율 확인", style='List Bullet')
    
    doc.save(str(RESEARCH_NOTES_DIR / "연구노트_전고체배터리_고체전해질_최적화.docx"))
    print("[OK] Created: 연구노트_전고체배터리_고체전해질_최적화.docx")


def create_bioplastic_note():
    """연구노트 2: 바이오 플라스틱 젖산 생산 수율 향상"""
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)
    
    # 제목
    h = doc.add_heading("바이오 플라스틱 합성을 위한 유산균 유래 유기산 생산 수율 향상 연구", level=1)
    h.runs[0].font.name = '맑은 고딕'
    h.runs[0].font.size = Pt(18)
    h.runs[0].font.bold = True
    
    # 메타데이터
    p_meta = doc.add_paragraph()
    p_meta.add_run("작성자: ").bold = True
    p_meta.add_run("이지현 선임연구원 (친환경바이오화학 연구센터)\n")
    p_meta.add_run("작성일: ").bold = True
    p_meta.add_run("2026-05-20\n")
    p_meta.add_run("프로젝트코드: ").bold = True
    p_meta.add_run("BIO-PL-2026-02\n")
    p_meta.add_run("문서보안등급: ").bold = True
    p_meta.add_run("일반 (Public)")
    
    # 1페이지: 서론 및 연구 배경
    doc.add_heading("1. 서론 및 연구 배경", level=2)
    doc.add_paragraph(
        "최근 탄소 배출 규제와 플라스틱 폐기물 문제 해결을 위해 생분해성 바이오 플라스틱인 PLA(Poly Lactic Acid)에 대한 "
        "관심이 크게 증가하고 있다. PLA 생산 원가의 핵심은 원료 물질인 L-Lactic acid(젖산)를 고순도, 고수율로 "
        "생산하는 발효 공정 기술이다."
    )
    doc.add_paragraph(
        "본 연구에서는 당사에서 분리한 유산균 변이균주인 Lactobacillus paracasei L-105의 배양 특성을 분석하고, "
        "탄소원 종류에 따른 젖산 생산 수율 및 생산 속도를 비교 평가하여 상용 배양 조건의 최적 인자를 규명하고자 한다."
    )
    
    doc.add_page_break()
    
    # 2페이지: 배양 조건 및 실험 설계
    doc.add_heading("2. 배양 조건 및 실험 설계", level=2)
    doc.add_paragraph(
        "실험은 두 가지 주요 탄소원인 포도당(Glucose)과 글리세롤(Glycerol)을 대상으로 병렬 배양을 진행하였다. "
        "발효 배지 조성은 MRS 배지를 기반으로 하였으며, 질소원으로 효모 추출물(Yeast Extract) 15 g/L 및 펩톤(Peptone) 10 g/L를 첨가하였다."
    )
    doc.add_paragraph("배양 상세 조건은 다음과 같다:")
    doc.add_paragraph("• 장비: 5L 교반 발효조 (5L Fermentor, Working Volume 3L)", style='List Bullet')
    doc.add_paragraph("• 배양 온도: 37.0°C 유지", style='List Bullet')
    doc.add_paragraph("• 교반 속도: 200 rpm (공기 주입량 1.0 vvm)", style='List Bullet')
    doc.add_paragraph("• pH 제어: pH 6.0 유지 (2 M NaOH 및 2 M HCl 용액을 이용한 자동 제어)", style='List Bullet')
    
    doc.add_paragraph(
        "배양 중 6시간 간격으로 샘플링을 실시하여 세포 농도(OD600)를 측정하고, "
        "원심분리된 상등액의 잔류 당 농도 및 생성된 젖산 농도를 HPLC(High-Performance Liquid Chromatography)로 정량 분석하였다."
    )
    
    doc.add_page_break()
    
    # 3페이지: 결과 및 고찰
    doc.add_heading("3. 결과 및 고찰", level=2)
    doc.add_paragraph(
        "포도당을 탄소원으로 사용한 경우, 배양 시작 36시간 만에 포도당이 완전히 소모되었으며, "
        "최종 젖산 축적량은 98 g/L를 달성하였다. 이때 젖산 수율(Yield)은 투입된 포도당 대비 0.94 g/g으로 나타났다."
    )
    doc.add_paragraph(
        "반면, 글리세롤을 탄소원으로 사용한 경우에는 균주의 대사 속도가 현저히 느려 72시간 경과 시점에도 글리세롤 소모가 완료되지 않았으며, "
        "최종 젖산 농도는 45 g/L, 수율은 0.62 g/g 수준에 그쳤다."
    )
    
    # 표 2 추가
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Table Grid'
    
    # 헤더
    hdr = table.rows[0].cells
    hdr[0].text = '탄소원 종류'
    hdr[1].text = '배양 완료 시간 (h)'
    hdr[2].text = '최종 젖산 농도 (g/L)'
    hdr[3].text = '젖산 수율 (g/g-당)'
    
    # 포도당 데이터
    row1 = table.rows[1].cells
    row1[0].text = '포도당 (Glucose)'
    row1[1].text = '36'
    row1[2].text = '98.0'
    row1[3].text = '0.94'
    
    # 글리세롤 데이터
    row2 = table.rows[2].cells
    row2[0].text = '글리세롤 (Glycerol)'
    row2[1].text = '72 (미완료)'
    row2[2].text = '45.0'
    row2[3].text = '0.62'
    
    doc.add_paragraph(
        "\n위 결과로부터 L. paracasei L-105 균주가 글리세롤보다 포도당 배지에서 훨씬 높은 대사 활성을 가짐을 확인하였다. "
        "특히 포도당 배지에서의 0.94 g/g 수율은 상용 플랜트 운전 시의 경제성 임계치인 0.90 g/g을 뛰어넘는 고효율이다."
    )
    
    doc.add_page_break()
    
    # 4페이지: 결론 및 향후 과제
    doc.add_heading("4. 결론 및 향후 과제", level=2)
    doc.add_paragraph(
        "1. L. paracasei L-105 균주의 젖산 발효에 적합한 탄소원은 포도당임을 재확인하였다. "
        "최적 조건에서 36시간 배양을 통해 최종 젖산 농도 98 g/L, 수율 0.94 g/g를 얻었다."
    )
    doc.add_paragraph(
        "2. 산업적 생산성 확대를 위해, 저렴한 폐당밀(Blackstrap Molasses)을 당원으로 대체 적용하는 배지 고도화 연구를 제안한다."
    )
    
    doc.save(str(RESEARCH_NOTES_DIR / "연구노트_바이오플라스틱_젖산생산_수율향상.docx"))
    print("[OK] Created: 연구노트_바이오플라스틱_젖산생산_수율향상.docx")


def create_solar_hydrogen_note():
    """연구노트 3: 페로브스카이트 광전극 표면 개질"""
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)
    
    # 제목
    h = doc.add_heading("태양광 수소 생산을 위한 페로브스카이트 광전극 표면 개질 및 광전환 효율 평가", level=1)
    h.runs[0].font.name = '맑은 고딕'
    h.runs[0].font.size = Pt(18)
    h.runs[0].font.bold = True
    
    # 메타데이터
    p_meta = doc.add_paragraph()
    p_meta.add_run("작성자: ").bold = True
    p_meta.add_run("박준서 책임연구원 (신재생에너지 소재센터)\n")
    p_meta.add_run("작성일: ").bold = True
    p_meta.add_run("2026-05-28\n")
    p_meta.add_run("프로젝트코드: ").bold = True
    p_meta.add_run("PV-H2-2026-07\n")
    p_meta.add_run("문서보안등급: ").bold = True
    p_meta.add_run("대외비 (Confidential)")
    
    # 1페이지: 연구 개요 및 배경
    doc.add_heading("1. 연구 개요 및 배경", level=2)
    doc.add_paragraph(
        "광전기화학적(PEC) 물분해를 통한 수소 생산 기술은 온실가스 배출 없이 청정 수소를 제조할 수 있는 미래 기술이다. "
        "페로브스카이트(Perovskite) 광전극은 우수한 광흡수 특성과 저렴한 제조 비용으로 주목받고 있으나, "
        "수용액 전해질 환경에서 수분 침투로 인해 유기-무기 하이브리드 페로브스카이트 소재가 매우 급격히 용해되거나 열화되는 문제를 가진다."
    )
    doc.add_paragraph(
        "본 연구에서는 페로브스카이트 광전극 표면에 원자층 증착법(Atomic Layer Deposition, ALD)을 적용하여 "
        "치밀한 Al2O3 보호막을 증착하고, 두께에 따른 광전류 밀도 변화와 수중 장기 내구성을 시험하여 최적 코팅 두께를 제안하고자 한다."
    )
    
    doc.add_page_break()
    
    # 2페이지: 시편 제작 공정
    doc.add_heading("2. 시편 제작 공정", level=2)
    doc.add_paragraph("실험에 사용된 페로브스카이트(MAPbI3) 광전극 시편은 다음과 같은 순서로 제작되었다:")
    
    doc.add_paragraph("• 기판 세정: FTO 투명 전극 유리 기판을 아세톤, 이소프로필알코올, 탈이온수 순으로 초음파 세척한다.", style='List Bullet')
    doc.add_paragraph("• 전자수송층(ETL) 형성: TiO2 페이스트를 스핀 코팅(3000 rpm, 30초)한 뒤, 500°C에서 30분간 열처리한다.", style='List Bullet')
    doc.add_paragraph("• 페로브스카이트 적층: MAPbI3 전구체 용액을 스핀 코팅하고, 100°C 핫플레이트 상에서 10분 동안 어닐링한다.", style='List Bullet')
    doc.add_paragraph(
        "• ALD 보호막 증착: 원자층 증착(ALD) 챔버 내부 온도 150°C 조건 하에서, 트리메틸알루미늄(TMA)과 수증기(H2O)를 교대로 "
        "주입하여 Al2O3 초박막을 증착한다. 증착 두께 조건을 각각 5 nm, 10 nm, 20 nm로 다양화하여 샘플을 제작하였다."
    )
    
    doc.add_page_break()
    
    # 3페이지: 광전기화학적 성능 측정 및 분석
    doc.add_heading("3. 광전기화학적 성능 측정 및 분석", level=2)
    doc.add_paragraph(
        "제작된 시편의 PEC 성능은 AM 1.5G 1 sun 인공태양 광원(100 mW/cm2) 하에서 평가되었다. "
        "측정에 사용된 전해질은 0.5 M Na2SO4 수용액(pH 6.8)이었으며, 기준전극은 Ag/AgCl, 대극은 백선(Pt wire)을 적용하였다."
    )
    doc.add_paragraph(
        "측정된 1.23 V vs. RHE(정상 가역 수소 전극 전위)에서의 광전류 밀도와 연속 광조사에 따른 성능 반감기(안정성) 결과는 "
        "아래 [표 3]과 같다."
    )
    
    # 표 3 추가
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    
    # 헤더
    hdr = table.rows[0].cells
    hdr[0].text = 'Al2O3 두께 (nm)'
    hdr[1].text = '광전류 밀도 (mA/cm2)'
    hdr[2].text = '개시 전위 (V vs. RHE)'
    hdr[3].text = '안정성 (24시간 후 유지율)'
    
    # 데이터 행 1 (보호막 없음)
    row1 = table.rows[1].cells
    row1[0].text = '0 nm (미적용)'
    row1[1].text = '12.2'
    row1[2].text = '0.45'
    row1[3].text = '1% 미만 (1시간 내 소멸)'
    
    # 데이터 행 2 (5 nm)
    row2 = table.rows[2].cells
    row2[0].text = '5 nm'
    row2[1].text = '15.6'
    row2[2].text = '0.38'
    row2[3].text = '42%'
    
    # 데이터 행 3 (10 nm)
    row3 = table.rows[3].cells
    row3[0].text = '10 nm (최적)'
    row3[1].text = '18.5'
    row3[2].text = '0.31'
    row3[3].text = '88%'
    
    # 데이터 행 4 (20 nm)
    row4 = table.rows[4].cells
    row4[0].text = '20 nm'
    row4[1].text = '8.7'
    row4[2].text = '0.52'
    row4[3].text = '95%'
    
    doc.add_paragraph(
        "\n실험 결과, Al2O3 보호막 두께가 10 nm일 때 18.5 mA/cm2의 최고 광전류 밀도를 나타냈으며, "
        "24시간 장기 조사 후에도 성능이 초기 대비 88% 수준으로 우수하게 유지되었다. "
        "두께가 20 nm로 과도하게 두꺼워지면 전도성 저하 및 터널링 저항 증가로 인해 광전류가 8.7 mA/cm2로 오히려 감소하였다."
    )
    
    doc.add_page_break()
    
    # 4페이지: 종합 결론
    doc.add_heading("4. 종합 결론", level=2)
    doc.add_paragraph(
        "ALD 공정을 통한 10 nm 두께의 Al2O3 보호층 증착을 통해 페로브스카이트 광전극의 수중 용해 문제를 해결하고, "
        "동시에 표면 결함 결합(Passivation) 효과를 유도하여 효율을 약 50% 향상시키는 시너지 효과를 달성하였다."
    )
    doc.add_paragraph(
        "향후 실용화를 위해 대면적(100 cm2 이상) 광전극에서의 두께 균일성(Uniformity) 분석과 "
        "안정성이 더욱 높은 NiOx 등 정공수송층과의 하이브리드 보호 시스템 설계에 대한 추가 검토를 진행하고자 한다."
    )
    
    doc.save(str(RESEARCH_NOTES_DIR / "연구노트_페로브스카이트_광전극_표면개질.docx"))
    print("[OK] Created: 연구노트_페로브스카이트_광전극_표면개질.docx")


def main():
    print("샘플 연구노트 생성 스크립트 실행 시작...")
    create_battery_note()
    create_bioplastic_note()
    create_solar_hydrogen_note()
    print("모든 샘플 연구노트가 정상적으로 생성되었습니다!")


if __name__ == "__main__":
    main()
